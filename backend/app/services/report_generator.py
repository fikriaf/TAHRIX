"""Report generator — PDF, Markdown, and DOCX from a LaTeX-inspired structured template.

Design: all content lives in tables (LaTeX tabular style) for a clean forensic look.
The HTML template mirrors a LaTeX document structure with full-width tables, ruled
separators, monospace addresses, and a header block — then WeasyPrint converts to PDF.
The same data is also serializable to Markdown and DOCX from the same source dict.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from jinja2 import Environment, select_autoescape

from app.core.logging import get_logger

logger = get_logger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# HTML/PDF template — LaTeX-inspired, all content in tables
# ─────────────────────────────────────────────────────────────────────────────
_HTML_TEMPLATE = r"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>TAHRIX Forensic Report — {{ case_id }}</title>
<style>
  @page { size: A4; margin: 22mm 18mm 22mm 18mm; }
  * { box-sizing: border-box; }
  body { font-family: 'Helvetica Neue', Arial, sans-serif; font-size: 9.5pt; color: #0f172a; margin: 0; line-height: 1.45; }

  /* ── HEADER ── */
  .doc-header { border-bottom: 2.5pt solid #0f172a; padding-bottom: 8pt; margin-bottom: 12pt; display: flex; justify-content: space-between; align-items: flex-end; }
  .doc-title { font-size: 18pt; font-weight: 700; letter-spacing: 1px; color: #0f172a; }
  .doc-sub { font-size: 8pt; color: #475569; letter-spacing: 3px; text-transform: uppercase; margin-top: 2pt; }
  .doc-meta-right { text-align: right; font-size: 8pt; color: #475569; }
  .doc-meta-right strong { color: #0f172a; font-size: 9pt; }

  /* ── SECTION TITLES ── */
  .sec-title { font-size: 9pt; font-weight: 700; letter-spacing: 3px; text-transform: uppercase; color: #0f172a; border-bottom: 1pt solid #0f172a; padding-bottom: 2pt; margin: 14pt 0 6pt; }

  /* ── TABLES ── */
  table { width: 100%; border-collapse: collapse; font-size: 9pt; margin-bottom: 4pt; }
  thead tr { background: #0f172a; color: #f8fafc; }
  thead th { padding: 5pt 7pt; text-align: left; font-weight: 600; font-size: 8pt; letter-spacing: 1px; text-transform: uppercase; }
  tbody tr { border-bottom: 0.5pt solid #e2e8f0; }
  tbody tr:nth-child(even) { background: #f8fafc; }
  tbody td { padding: 4.5pt 7pt; vertical-align: top; }
  td.label { font-weight: 600; color: #334155; width: 32%; white-space: nowrap; }
  td.value { color: #0f172a; word-break: break-all; }

  /* ── VERDICT BANNER ── */
  .verdict-banner { padding: 10pt 14pt; margin: 6pt 0 10pt; display: flex; align-items: center; gap: 18pt; }
  .verdict-banner.critical { background: #fef2f2; border-left: 4pt solid #dc2626; }
  .verdict-banner.high     { background: #fff7ed; border-left: 4pt solid #ea580c; }
  .verdict-banner.medium   { background: #fefce8; border-left: 4pt solid #ca8a04; }
  .verdict-banner.low      { background: #f0fdf4; border-left: 4pt solid #16a34a; }
  .verdict-score { font-size: 32pt; font-weight: 700; line-height: 1; }
  .verdict-banner.critical .verdict-score { color: #dc2626; }
  .verdict-banner.high     .verdict-score { color: #ea580c; }
  .verdict-banner.medium   .verdict-score { color: #ca8a04; }
  .verdict-banner.low      .verdict-score { color: #16a34a; }
  .verdict-detail { flex: 1; }
  .verdict-grade { font-size: 11pt; font-weight: 700; letter-spacing: 3px; text-transform: uppercase; }
  .verdict-expl { font-size: 8pt; color: #475569; margin-top: 3pt; line-height: 1.5; max-height: 80pt; overflow: hidden; }

  /* ── ANOMALY FLAGS ── */
  .anom-on { color: #dc2626; font-weight: 700; }
  .anom-off { color: #94a3b8; }

  /* ── RISK BAR ── */
  .rbar-wrap { display: flex; align-items: center; gap: 6pt; }
  .rbar-bg { flex: 1; height: 5pt; background: #e2e8f0; border-radius: 2pt; overflow: hidden; }
  .rbar-fill { height: 100%; background: #0f172a; border-radius: 2pt; }
  .rbar-val { font-size: 8pt; color: #475569; width: 28pt; text-align: right; }

  /* ── CODE / MONO ── */
  code { font-family: 'Courier New', Courier, monospace; font-size: 8pt; background: #f1f5f9; padding: 1pt 3pt; border-radius: 2pt; }
  .mono { font-family: 'Courier New', Courier, monospace; font-size: 8pt; }

  /* ── THREAT PILL ── */
  .threat-high { color: #dc2626; font-weight: 700; }
  .threat-med  { color: #ea580c; font-weight: 600; }
  .threat-low  { color: #ca8a04; }

  /* ── FOOTER ── */
  .doc-footer { margin-top: 18pt; padding-top: 6pt; border-top: 1pt solid #cbd5e1; font-size: 7.5pt; color: #94a3b8; display: flex; justify-content: space-between; }
</style>
</head>
<body>

<!-- ══ HEADER ══════════════════════════════════════════════════════════ -->
<div class="doc-header">
  <div>
    <div class="doc-title">TAHRIX</div>
    <div class="doc-sub">Agentic Blockchain Intelligence · Forensic Report</div>
  </div>
  <div class="doc-meta-right">
    <div><strong>{{ case_id }}</strong></div>
    <div>Generated: {{ generated_at }} UTC</div>
    <div>Classification: RESTRICTED</div>
  </div>
</div>

<!-- ══ VERDICT BANNER ════════════════════════════════════════════════════ -->
<div class="verdict-banner {{ risk.grade }}">
  <div class="verdict-score">{{ '%.0f'|format(risk.score) }}</div>
  <div class="verdict-detail">
    <div class="verdict-grade">{{ risk.grade }} Risk</div>
    <div style="font-size:10pt;color:#64748b;margin-top:1pt">Confidence: {{ '%.0f'|format((agent.confidence or 0) * 100) }}%</div>
  </div>
</div>

<!-- ══ SECTION 1: SUBJECT ════════════════════════════════════════════════ -->
<div class="sec-title">1 · Subject of Investigation</div>
<table><tbody>
  <tr><td class="label">Target Address</td><td class="value"><code>{{ address }}</code></td></tr>
  <tr><td class="label">Blockchain Network</td><td class="value">{{ chain }}</td></tr>
  <tr><td class="label">Case Reference</td><td class="value"><code>{{ case_id }}</code></td></tr>
  <tr><td class="label">Analysis Depth</td><td class="value">{{ agent.depth or 3 }} hops</td></tr>
  <tr><td class="label">Investigation Started</td><td class="value">{{ agent.started_at or '—' }}</td></tr>
  <tr><td class="label">Investigation Completed</td><td class="value">{{ agent.completed_at or '—' }}</td></tr>
  {% if agent.duration_s is not none %}<tr><td class="label">Duration</td><td class="value">{{ agent.duration_s }}s</td></tr>{% endif %}
  <tr><td class="label">OFAC Sanctions</td><td class="value">{% if risk.sanctions and risk.sanctions.sanctioned %}<strong style="color:#dc2626">SANCTIONED — OFAC SDN LIST MATCH</strong>{% else %}No match found{% endif %}</td></tr>
  <tr><td class="label">GNN Illicit Probability</td><td class="value">{% if risk.gnn %}{{ '%.4f'|format(risk.gnn.score) }} — {{ risk.gnn.label.value }}{% else %}Unavailable{% endif %}</td></tr>
  {% if agent.ipfs_cid %}<tr><td class="label">IPFS Evidence Hash</td><td class="value"><code>{{ agent.ipfs_cid }}</code></td></tr>{% endif %}
</tbody></table>

<!-- ══ SECTION 2: RISK SCORING ══════════════════════════════════════════ -->
<div class="sec-title">2 · Risk Scoring Breakdown</div>
<table>
  <thead><tr><th>Factor</th><th>Raw (0–1)</th><th>Weight</th><th>Points</th><th>Bar</th></tr></thead>
  <tbody>
    <tr>
      <td class="label">GNN — Graph Neural Network</td>
      <td>{{ '%.4f'|format(risk.components.gnn) }}</td>
      <td>{{ '%.0f'|format(risk.components.w_gnn * 100) }}%</td>
      <td>{{ '%.1f'|format(risk.components.gnn * risk.components.w_gnn * 100) }}</td>
      <td><div class="rbar-wrap"><div class="rbar-bg"><div class="rbar-fill" style="width:{{ (risk.components.gnn * 100)|int }}%"></div></div><span class="rbar-val">{{ '%.0f'|format(risk.components.gnn * 100) }}%</span></div></td>
    </tr>
    <tr>
      <td class="label">Anomaly Patterns</td>
      <td>{{ '%.4f'|format(risk.components.anomaly) }}</td>
      <td>{{ '%.0f'|format(risk.components.w_anomaly * 100) }}%</td>
      <td>{{ '%.1f'|format(risk.components.anomaly * risk.components.w_anomaly * 100) }}</td>
      <td><div class="rbar-wrap"><div class="rbar-bg"><div class="rbar-fill" style="width:{{ (risk.components.anomaly * 100)|int }}%"></div></div><span class="rbar-val">{{ '%.0f'|format(risk.components.anomaly * 100) }}%</span></div></td>
    </tr>
    <tr>
      <td class="label">OFAC Sanctions</td>
      <td>{{ '%.4f'|format(risk.components.sanctions) }}</td>
      <td>{{ '%.0f'|format(risk.components.w_sanctions * 100) }}%</td>
      <td>{{ '%.1f'|format(risk.components.sanctions * risk.components.w_sanctions * 100) }}</td>
      <td><div class="rbar-wrap"><div class="rbar-bg"><div class="rbar-fill" style="width:{{ (risk.components.sanctions * 100)|int }}%"></div></div><span class="rbar-val">{{ '%.0f'|format(risk.components.sanctions * 100) }}%</span></div></td>
    </tr>
    <tr>
      <td class="label">Threat Intelligence</td>
      <td>{{ '%.4f'|format(risk.components.get('threat', 0)) }}</td>
      <td>{{ '%.0f'|format(risk.components.get('w_threat', 0) * 100) }}%</td>
      <td>{{ '%.1f'|format(risk.components.get('threat', 0) * risk.components.get('w_threat', 0) * 100) }}</td>
      <td><div class="rbar-wrap"><div class="rbar-bg"><div class="rbar-fill" style="width:{{ (risk.components.get('threat', 0) * 100)|int }}%"></div></div><span class="rbar-val">{{ '%.0f'|format(risk.components.get('threat', 0) * 100) }}%</span></div></td>
    </tr>
    <tr>
      <td class="label">Network Centrality</td>
      <td>{{ '%.4f'|format(risk.components.centrality) }}</td>
      <td>{{ '%.0f'|format(risk.components.w_centrality * 100) }}%</td>
      <td>{{ '%.1f'|format(risk.components.centrality * risk.components.w_centrality * 100) }}</td>
      <td><div class="rbar-wrap"><div class="rbar-bg"><div class="rbar-fill" style="width:{{ (risk.components.centrality * 100)|int }}%"></div></div><span class="rbar-val">{{ '%.0f'|format(risk.components.centrality * 100) }}%</span></div></td>
    </tr>
    <tr style="background:#0f172a;color:#f8fafc;font-weight:700">
      <td colspan="3" style="padding:5pt 7pt;color:#f8fafc">COMPOSITE RISK SCORE</td>
      <td style="padding:5pt 7pt;color:#f8fafc;font-size:12pt">{{ '%.1f'|format(risk.score) }} / 100</td>
      <td></td>
    </tr>
  </tbody>
</table>

{% if is_osint %}
<!-- ══ SECTION 3: OSINT FINDINGS ═══════════════════════════════════════ -->
<div class="sec-title">3 · OSINT Findings</div>
<table><tbody>
  <tr><td class="label">Investigation Type</td><td class="value"><strong>Open-Source Intelligence</strong> — Non-wallet entity</td></tr>
  <tr><td class="label">Entity Name</td><td class="value"><code>{{ address }}</code></td></tr>
  <tr><td class="label">OSINT Sources Queried</td><td class="value">
    {% if agent.osint_sources %}
      {{ agent.osint_sources|join(', ') }}
    {% else %}
      Sherlock, TheHarvester, Web Search, Social Media Intel
    {% endif %}
  </td></tr>
  <tr><td class="label">Social Profiles Found</td><td class="value">{{ agent.social_found or 0 }}</td></tr>
  <tr><td class="label">Emails Found</td><td class="value">{{ agent.emails_found or 0 }}</td></tr>
  <tr><td class="label">Subdomains Found</td><td class="value">{{ agent.subdomains_found or 0 }}</td></tr>
  <tr><td class="label">Web Mentions</td><td class="value">{{ agent.osint_mentions or 0 }}</td></tr>
  <tr><td class="label">Threat Intel Matches</td><td class="value">{% if agent.threats %}{{ agent.threats|length }}{% else %}0{% endif %}</td></tr>
</tbody></table>

{% if agent.osint_details %}
<div class="sec-sub">Key OSINT Details</div>
<table><tbody>
{% for detail in agent.osint_details %}
  <tr><td class="label">{{ detail.source }}</td><td class="value">{{ detail.findings }}</td></tr>
{% endfor %}
</tbody></table>
{% endif %}

<!-- ══ SECTION 4: ON-CHAIN ACTIVITY (if wallet discovered) ═════════════════ -->
{% if agent.transactions_collected > 0 %}
<div class="sec-title">4 · On-Chain Activity (Wallet Linked)</div>
{% else %}
<div class="sec-title">4 · On-Chain Activity</div>
{% endif %}
{% else %}
<!-- ══ SECTION 3: ON-CHAIN ACTIVITY ═════════════════════════════════════ -->
{% endif %}
<div class="sec-title">3 · On-Chain Activity</div>
<table><tbody>
  <tr><td class="label">Transactions Collected</td><td class="value">{{ agent.transactions_collected or '—' }}</td></tr>
  <tr><td class="label">First Seen On-Chain</td><td class="value">{{ agent.first_seen or '—' }}</td></tr>
  <tr><td class="label">Last Active On-Chain</td><td class="value">{{ agent.last_seen or '—' }}</td></tr>
  <tr><td class="label">Forward Trace (outflow)</td><td class="value">{{ agent.trace_fwd_count or 0 }} counterparties</td></tr>
  <tr><td class="label">Backward Trace (inflow)</td><td class="value">{{ agent.trace_bwd_count or 0 }} counterparties</td></tr>
  <tr><td class="label">GNN Subgraph Size</td><td class="value">{{ agent.gnn_subgraph_size or '—' }} nodes processed</td></tr>
</tbody></table>

<!-- ══ SECTION 4: GRAPH TOPOLOGY ════════════════════════════════════════ -->
<div class="sec-title">4 · Knowledge Graph Topology</div>
<table>
  <thead><tr><th>Node Type</th><th>Count</th></tr></thead>
  <tbody>
    {% for ntype, count in agent.node_type_counts.items() %}
    <tr><td class="label">{{ ntype }}</td><td>{{ count }}</td></tr>
    {% else %}
    <tr><td colspan="2" style="color:#94a3b8">No graph data available</td></tr>
    {% endfor %}
    <tr style="font-weight:700;background:#f1f5f9">
      <td class="label">Total Nodes</td><td>{{ agent.graph_node_count }}</td>
    </tr>
    <tr style="font-weight:700;background:#f1f5f9">
      <td class="label">Total Edges</td><td>{{ agent.graph_edge_count }}</td>
    </tr>
  </tbody>
</table>

{% if agent.top_counterparties %}
<div style="margin-top:6pt;font-size:8.5pt;font-weight:700;color:#334155;letter-spacing:1px">TOP COUNTERPARTIES BY TX COUNT</div>
<table style="margin-top:4pt">
  <thead><tr><th>#</th><th>Address</th><th>Chain</th><th>TX Count</th><th>Balance (USD)</th><th>Flags</th></tr></thead>
  <tbody>
    {% for n in agent.top_counterparties %}
    <tr>
      <td>{{ loop.index }}</td>
      <td><code>{{ (n.address or n.id or '')[:20] }}…</code></td>
      <td>{{ n.chain or '—' }}</td>
      <td>{{ n.tx_count or '—' }}</td>
      <td>{{ '$%.2f'|format(n.balance_usd) if n.balance_usd else '—' }}</td>
      <td>{% if n.is_sanctioned %}<span class="threat-high">SANCTIONED</span>{% else %}—{% endif %}</td>
    </tr>
    {% endfor %}
  </tbody>
</table>
{% endif %}

<!-- ══ SECTION 5: THREAT INTELLIGENCE ══════════════════════════════════ -->
{% if agent.threats %}
<div class="sec-title">5 · Threat Intelligence Hits</div>
<table>
  <thead><tr><th>Source</th><th>Type</th><th>Name</th><th>Severity</th><th>Description</th></tr></thead>
  <tbody>
    {% for t in agent.threats %}
    <tr>
      <td class="mono">{{ t.source or '—' }}</td>
      <td>{{ t.type or t.get('threat_type','—') }}</td>
      <td style="font-weight:600">{{ t.name or '—' }}</td>
      <td>
        {% set sev = (t.severity or 0) * 100 %}
        <span class="{% if sev >= 80 %}threat-high{% elif sev >= 50 %}threat-med{% else %}threat-low{% endif %}">
          {{ '%.0f'|format(sev) }}%
        </span>
      </td>
      <td style="font-size:8pt;color:#475569">{{ (t.description or '—')[:120] }}</td>
    </tr>
    {% endfor %}
  </tbody>
</table>
{% endif %}

<!-- ══ SECTION 6: ANOMALY PATTERN DETECTION ═════════════════════════════ -->
<div class="sec-title">6 · Anomaly Pattern Detection (P01–P17)</div>
<table>
  <thead><tr><th>Code</th><th>Pattern</th><th>Status</th><th>Severity</th><th>Description</th></tr></thead>
  <tbody>
    {% set triggered_codes = risk.anomaly_flags | map(attribute='code') | map('string') | list %}
    {% for code, name in anomaly_defs %}
    {% set triggered = code in triggered_codes %}
    <tr>
      <td><code>{{ code }}</code></td>
      <td class="{% if triggered %}anom-on{% else %}anom-off{% endif %}">{{ name }}</td>
      <td>{% if triggered %}<span class="anom-on">● TRIGGERED</span>{% else %}<span class="anom-off">○ clear</span>{% endif %}</td>
      <td>
        {% if triggered %}
          {% for f in risk.anomaly_flags %}{% if f.code|string == code %}{{ '%.0f'|format(f.severity * 100) }}%{% endif %}{% endfor %}
        {% else %}—{% endif %}
      </td>
      <td style="font-size:8pt;color:#475569">
        {% if triggered %}
          {% for f in risk.anomaly_flags %}{% if f.code|string == code %}{{ f.description }}{% endif %}{% endfor %}
        {% else %}—{% endif %}
      </td>
    </tr>
    {% endfor %}
  </tbody>
</table>

<!-- ══ SECTION 7: GNN ANALYSIS ══════════════════════════════════════════ -->
{% if risk.gnn %}
<div class="sec-title">7 · Graph Neural Network Analysis</div>
<table><tbody>
  <tr><td class="label">Model</td><td class="value">GAT (Graph Attention Network) — Elliptic AUC 0.9684</td></tr>
  <tr><td class="label">Illicit Probability</td><td class="value">{{ '%.6f'|format(risk.gnn.score) }}</td></tr>
  <tr><td class="label">Classification</td><td class="value" style="font-weight:700">{{ risk.gnn.label.value }}</td></tr>
  <tr><td class="label">Subgraph Size</td><td class="value">{{ agent.gnn_subgraph_size or '—' }} nodes</td></tr>
  <tr><td class="label">Model Interpretation</td><td class="value" style="font-size:8.5pt">{{ risk.gnn.explanation or '—' }}</td></tr>
</tbody></table>
{% if risk.gnn.shap_top_features %}
<div style="margin-top:6pt;font-size:8.5pt;font-weight:700;color:#334155;letter-spacing:1px">SHAP FEATURE ATTRIBUTION</div>
<table style="margin-top:4pt">
  <thead><tr><th>#</th><th>Feature</th><th>Raw Value</th><th>SHAP Δ</th><th>Influence</th></tr></thead>
  <tbody>
    {% for feat in risk.gnn.shap_top_features[:10] %}
    <tr>
      <td>{{ loop.index }}</td>
      <td class="label">{{ feat.feature }}</td>
      <td class="mono">{{ '%.2f'|format(feat.raw_value) if feat.raw_value is not none else '—' }}</td>
      <td class="mono">{{ '%+.4f'|format(feat.value) }}</td>
      <td>
        <div class="rbar-wrap">
          <div class="rbar-bg"><div class="rbar-fill" style="width:{{ [((feat.value | abs) * 200)|int, 100] | min }}%"></div></div>
          <span class="rbar-val">{{ '%+.3f'|format(feat.value) }}</span>
        </div>
      </td>
    </tr>
    {% endfor %}
  </tbody>
</table>
{% endif %}
{% endif %}

<!-- ══ SECTION 8: INVESTIGATION EXECUTION ═══════════════════════════════ -->
<div class="sec-title">8 · Investigation Execution</div>
<table><tbody>
  <tr><td class="label">Cognitive Iterations</td><td class="value">{{ agent.iterations }}</td></tr>
  <tr><td class="label">Agent Confidence</td><td class="value">{{ '%.1f'|format((agent.confidence or 0) * 100) }}%</td></tr>
  <tr><td class="label">Duration</td><td class="value">{% if agent.duration_s %}{{ agent.duration_s }}s{% else %}—{% endif %}</td></tr>
  {% if agent.llm_usage %}<tr><td class="label">LLM Tokens Used</td><td class="value">{{ agent.llm_usage.get('total_tokens', '—') }} (prompt: {{ agent.llm_usage.get('prompt_tokens','—') }}, completion: {{ agent.llm_usage.get('completion_tokens','—') }})</td></tr>{% endif %}
  <tr><td class="label">Transactions Analyzed</td><td class="value">{{ agent.transactions_collected or 0 }}</td></tr>
  <tr><td class="label">Forward Counterparties</td><td class="value">{{ agent.trace_fwd_count or 0 }}</td></tr>
  <tr><td class="label">Backward Counterparties</td><td class="value">{{ agent.trace_bwd_count or 0 }}</td></tr>
  <tr><td class="label">Bridge Events</td><td class="value">{{ agent.bridge_events | length }}</td></tr>
</tbody></table>

<!-- ══ SECTION 9: AGENT VERDICT ═════════════════════════════════════════ -->
<div class="sec-title">9 · Agent Final Verdict</div>
<div style="background:#f8fafc;border-left:3pt solid #cbd5e1;padding:8pt 10pt;font-size:8.5pt;line-height:1.6;color:#0f172a">
  {{ agent.final_text or '— Agent did not produce a final text verdict.' }}
</div>

<!-- ══ FOOTER ════════════════════════════════════════════════════════════ -->
<div class="doc-footer">
  <span>TAHRIX Agentic AI Cyber Intelligence Platform · Restricted Distribution</span>
  <span>Risk scores are advisory; verify with a qualified analyst before legal action.</span>
</div>
</body>
</html>
"""
_ANOMALY_DEFS = [
    ("P01", "MIXER — Direct interaction with known mixer"),
    ("P02", "LAYERING — Multi-hop near-equal value transfers"),
    ("P03", "FAN_OUT — Rapid dispersal to many wallets"),
    ("P04", "FAN_IN — Rapid aggregation from many wallets"),
    ("P05", "PEELING — Sequential small outflows"),
    ("P06", "ROUND_TRIP — Funds returned to origin"),
    ("P07", "BRIDGE_HOP — Cross-chain bridge abuse"),
    ("P08", "WHALE — Unusually large transaction"),
    ("P09", "DORMANT — Long-dormant address suddenly active"),
    ("P10", "RAPID — High-frequency micro-transactions"),
    ("P11", "OFAC_IND — Indirect OFAC sanctioned exposure"),
    ("P12", "DEX_WASH — DEX wash-trading pattern"),
    ("P13", "NFT_WASH — NFT wash-trading pattern"),
    ("P14", "FLASH_LOAN — Flash loan attack pattern"),
    ("P15", "ADDR_POISON — Address poisoning attack"),
    ("P16", "RUG_PULL — Rug pull pattern"),
    ("P17", "SANDWICH — Sandwich attack MEV pattern"),
]

_env = Environment(autoescape=select_autoescape(["html", "xml"]))


# ─────────────────────────────────────────────────────────────────────────────
# Graph SVG renderer (pure Python — no D3, no browser)
# ─────────────────────────────────────────────────────────────────────────────
_NODE_COLOR = {
    "Wallet":        "#3b82f6",
    "Transaction":   "#a855f7",
    "Entity":        "#06b6d4",
    "OsintNode":     "#f59e0b",
    "ThreatIntel":   "#ef4444",
    "AnomalyPattern":"#f97316",
    "ChainNetwork":  "#10b981",
}
_NODE_W = {"Wallet": 120, "Transaction": 100, "Entity": 110,
           "OsintNode": 110, "ThreatIntel": 110,
           "AnomalyPattern": 100, "ChainNetwork": 90}
_NODE_H = {"Wallet": 40, "Transaction": 32, "Entity": 36,
           "OsintNode": 36, "ThreatIntel": 36,
           "AnomalyPattern": 32, "ChainNetwork": 30}


def _node_label(n: dict) -> str:
    t = n.get("node_type") or n.get("type", "Wallet")
    if t == "Wallet":      return (n.get("address") or n.get("id", ""))[:14] + "…"
    if t == "Transaction": return (n.get("hash") or n.get("id", ""))[:13] + "…"
    if t == "Entity":      return (n.get("name") or n.get("id", ""))[:16]
    if t == "OsintNode":   return (n.get("platform") or "web") + " · " + (n.get("source") or "")[:10]
    if t == "ThreatIntel": return (n.get("threat_type") or n.get("source") or "")[:16]
    if t == "AnomalyPattern": return n.get("code") or "PATTERN"
    if t == "ChainNetwork": return (n.get("name") or n.get("id", ""))[:12]
    return (n.get("id") or "")[:14]


def _layout_nodes(nodes: list[dict], edges: list[dict],
                  width: int = 900, height: int = 500) -> dict[str, tuple[float, float]]:
    """Simple iterative force-directed layout (repulsion + spring).
    Runs ~60 iterations — fast enough for a report render."""
    import math, random
    random.seed(42)
    pos: dict[str, list[float]] = {}
    ids = [n["id"] for n in nodes]

    # Seed positions on a circle
    n = len(ids)
    cx, cy = width / 2, height / 2
    r = min(width, height) * 0.38
    for i, nid in enumerate(ids):
        angle = 2 * math.pi * i / max(n, 1)
        pos[nid] = [cx + r * math.cos(angle), cy + r * math.sin(angle)]

    # Build adjacency for spring force
    adj: dict[str, list[str]] = {nid: [] for nid in ids}
    for e in edges:
        s, t = e.get("source") or e.get("from", ""), e.get("target") or e.get("to", "")
        if s in adj: adj[s].append(t)
        if t in adj: adj[t].append(s)

    k = math.sqrt(width * height / max(n, 1))

    for iteration in range(80):
        disp: dict[str, list[float]] = {nid: [0.0, 0.0] for nid in ids}
        temp = 30 * (1 - iteration / 80)  # cooling

        # Repulsion between all pairs
        for i, u in enumerate(ids):
            for v in ids[i+1:]:
                dx = pos[u][0] - pos[v][0]
                dy = pos[u][1] - pos[v][1]
                dist = math.hypot(dx, dy) or 0.1
                force = k * k / dist
                fx, fy = force * dx / dist, force * dy / dist
                disp[u][0] += fx; disp[u][1] += fy
                disp[v][0] -= fx; disp[v][1] -= fy

        # Attraction along edges
        for e in edges:
            s, t = e.get("source") or e.get("from", ""), e.get("target") or e.get("to", "")
            if s not in pos or t not in pos:
                continue
            dx = pos[s][0] - pos[t][0]
            dy = pos[s][1] - pos[t][1]
            dist = math.hypot(dx, dy) or 0.1
            force = dist * dist / k
            fx, fy = force * dx / dist, force * dy / dist
            disp[s][0] -= fx; disp[s][1] -= fy
            disp[t][0] += fx; disp[t][1] += fy

        # Apply displacement with temperature cap
        for nid in ids:
            d = disp[nid]
            dist = math.hypot(*d) or 0.1
            move = min(dist, temp)
            pos[nid][0] += d[0] / dist * move
            pos[nid][1] += d[1] / dist * move
            # Clamp to canvas
            pos[nid][0] = max(80, min(width  - 80, pos[nid][0]))
            pos[nid][1] = max(40, min(height - 40, pos[nid][1]))

    return {nid: (pos[nid][0], pos[nid][1]) for nid in ids}


def render_graph_svg(graph_data: dict | None, width: int = 900, height: int = 500) -> str:
    """Render graph nodes/edges to an inline SVG string for the PDF report.
    Returns empty string if no data or empty graph."""
    if not graph_data:
        return ""

    raw_nodes: list[dict] = graph_data.get("nodes") or []
    raw_edges: list[dict] = graph_data.get("edges") or []
    if not raw_nodes:
        return ""

    # Normalize node type
    for n in raw_nodes:
        n["type"] = n.get("node_type") or n.get("type") or "Wallet"
        n["id"] = str(n.get("id") or n.get("address") or "")

    # Normalize edge source/target
    for e in raw_edges:
        e["source"] = str(e.get("source") or e.get("from") or "")
        e["target"] = str(e.get("target") or e.get("to") or "")

    # Filter isolated nodes (same as frontend)
    connected: set[str] = set()
    for e in raw_edges:
        connected.add(e["source"]); connected.add(e["target"])
    nodes = [n for n in raw_nodes if n["id"] in connected or n.get("is_focal")]
    node_ids = {n["id"] for n in nodes}
    edges = [e for e in raw_edges if e["source"] in node_ids and e["target"] in node_ids]

    if not nodes:
        return ""

    pos = _layout_nodes(nodes, edges, width, height)

    lines: list[str] = []
    lines.append(
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" '
        f'viewBox="0 0 {width} {height}" style="background:#060a0f;border-radius:4px">'
    )
    # Arrow marker
    lines.append(
        '<defs><marker id="arr" viewBox="0 -4 8 8" refX="8" refY="0" '
        'markerWidth="5" markerHeight="5" orient="auto">'
        '<path d="M0,-4L8,0L0,4" fill="#243552"/></marker></defs>'
    )

    # Draw edges first
    for e in edges:
        s, t = e["source"], e["target"]
        if s not in pos or t not in pos:
            continue
        x1, y1 = pos[s]; x2, y2 = pos[t]
        # Shorten line to not overlap node rects
        import math
        dx, dy = x2 - x1, y2 - y1
        dist = math.hypot(dx, dy) or 1
        shrink = 22
        sx1 = x1 + dx / dist * shrink; sy1 = y1 + dy / dist * shrink
        sx2 = x2 - dx / dist * shrink; sy2 = y2 - dy / dist * shrink
        label = e.get("edge_type") or e.get("type") or ""
        lines.append(
            f'<line x1="{sx1:.1f}" y1="{sy1:.1f}" x2="{sx2:.1f}" y2="{sy2:.1f}" '
            f'stroke="#243552" stroke-width="1" marker-end="url(#arr)"/>'
        )
        if label:
            mx, my = (sx1 + sx2) / 2, (sy1 + sy2) / 2
            lines.append(
                f'<text x="{mx:.1f}" y="{my:.1f}" fill="#334155" '
                f'font-size="7" font-family="monospace" text-anchor="middle" dy="-3">'
                f'{label[:12]}</text>'
            )

    # Draw nodes
    for n in nodes:
        nid = n["id"]
        if nid not in pos:
            continue
        cx, cy = pos[nid]
        t = n["type"]
        nw = _NODE_W.get(t, 100)
        nh = _NODE_H.get(t, 34)
        color = _NODE_COLOR.get(t, "#243552")
        is_focal = n.get("is_focal", False)
        is_sanctioned = n.get("is_sanctioned", False)
        stroke = "#ef4444" if is_sanctioned else ("#3b82f6" if is_focal else color)
        sw = 1.5 if (is_focal or is_sanctioned) else 1
        fill = "#0d1f3c" if is_focal else "#0c1118"
        x0, y0 = cx - nw / 2, cy - nh / 2

        # Background rect
        lines.append(
            f'<rect x="{x0:.1f}" y="{y0:.1f}" width="{nw}" height="{nh}" '
            f'rx="3" fill="{fill}" stroke="{stroke}" stroke-width="{sw}"/>'
        )
        # Type badge stripe
        lines.append(
            f'<rect x="{x0:.1f}" y="{y0:.1f}" width="{nw}" height="10" '
            f'rx="3" fill="{color}" opacity="0.18"/>'
        )
        # Type label
        type_lbl = {"Wallet": "WALLET", "Transaction": "TX", "Entity": "ENTITY",
                    "OsintNode": "OSINT", "ThreatIntel": "THREAT",
                    "AnomalyPattern": "ANOMALY", "ChainNetwork": "CHAIN"}.get(t, t[:6].upper())
        lines.append(
            f'<text x="{cx:.1f}" y="{y0 + 7:.1f}" fill="{color}" '
            f'font-size="7" font-family="monospace" text-anchor="middle" '
            f'letter-spacing="1.5">{type_lbl}</text>'
        )
        # Node label
        lbl = _node_label(n)
        lines.append(
            f'<text x="{cx:.1f}" y="{cy + 4:.1f}" fill="#94a3b8" '
            f'font-size="8" font-family="monospace" text-anchor="middle">{lbl}</text>'
        )

    lines.append("</svg>")
    return "\n".join(lines)


def _build_context(*, case_id: str, risk, agent_result: dict[str, Any],
                   graph_svg: str | None = None) -> dict[str, Any]:
    chain_val = risk.chain.value if hasattr(risk.chain, "value") else str(risk.chain)
    return dict(
        case_id=case_id,
        generated_at=datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S"),
        address=risk.address,
        chain=chain_val,
        is_osint=(chain_val == "OSINT"),
        risk=risk,
        agent=agent_result,
        anomaly_defs=_ANOMALY_DEFS,
        graph_svg=graph_svg or "",
    )


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────
def render_html(*, case_id: str, risk, agent_result: dict[str, Any],
                graph_svg: str | None = None) -> str:
    tpl = _env.from_string(_HTML_TEMPLATE)
    return tpl.render(**_build_context(case_id=case_id, risk=risk,
                                       agent_result=agent_result, graph_svg=graph_svg))


def build_pdf_report(*, case_id: str, risk, agent_result: dict[str, Any],
                     graph_svg: str | None = None) -> bytes:
    """Render HTML template → PDF via WeasyPrint."""
    html = render_html(case_id=case_id, risk=risk, agent_result=agent_result,
                       graph_svg=graph_svg)
    from weasyprint import HTML
    return HTML(string=html).write_pdf()


# ─────────────────────────────────────────────────────────────────────────────
# Markdown
# ─────────────────────────────────────────────────────────────────────────────
def build_markdown_report(*, case_id: str, risk, agent_result: dict[str, Any],
                          graph_svg: str | None = None) -> bytes:
    """Generate a structured Markdown report — 9 sections matching PDF."""
    ctx = _build_context(case_id=case_id, risk=risk, agent_result=agent_result,
                         graph_svg=graph_svg)
    grade = (risk.grade or "low").upper()
    score = risk.score or 0
    address = risk.address
    chain = ctx["chain"]
    now = ctx["generated_at"]
    expl = risk.explanation or "—"
    a = agent_result  # shorthand

    # ── Sanctions / GNN header fields
    sanctions_txt = "**SANCTIONED — OFAC SDN LIST MATCH**" if (
        risk.sanctions and risk.sanctions.sanctioned) else "No match found"
    gnn_txt = (f"{risk.gnn.score:.4f} ({risk.gnn.label.value})"
               if risk.gnn else "Unavailable")

    # ── Section 2: Risk component rows
    c = risk.components
    comp_rows = [
        f"| GNN | {c.get('gnn', 0):.3f} | {c.get('w_gnn', 0)*100:.0f}% | {c.get('gnn', 0)*c.get('w_gnn', 0)*100:.1f} |",
        f"| Anomaly | {c.get('anomaly', 0):.3f} | {c.get('w_anomaly', 0)*100:.0f}% | {c.get('anomaly', 0)*c.get('w_anomaly', 0)*100:.1f} |",
        f"| Sanctions | {c.get('sanctions', 0):.3f} | {c.get('w_sanctions', 0)*100:.0f}% | {c.get('sanctions', 0)*c.get('w_sanctions', 0)*100:.1f} |",
        f"| Threat Intel | {c.get('threat', 0):.3f} | {c.get('w_threat', 0)*100:.0f}% | {c.get('threat', 0)*c.get('w_threat', 0)*100:.1f} |",
        f"| Centrality | {c.get('centrality', 0):.3f} | {c.get('w_centrality', 0)*100:.0f}% | {c.get('centrality', 0)*c.get('w_centrality', 0)*100:.1f} |",
        f"| **COMPOSITE** | **{score:.1f} / 100** | | |",
    ]

    # ── Section 3: On-Chain Activity
    tx_info = a.get("tx_info") or {}
    onchain_rows = [
        f"| Transactions Analysed | {tx_info.get('count', a.get('transactions_collected', 0))} |",
        f"| First Seen | {tx_info.get('first_seen') or '—'} |",
        f"| Last Seen | {tx_info.get('last_seen') or '—'} |",
        f"| Forward Trace Hops | {a.get('trace_fwd_count', 0)} |",
        f"| Backward Trace Hops | {a.get('trace_bwd_count', 0)} |",
        f"| Bridge Events | {len(a.get('bridge_events', []))} |",
        f"| IPFS Archival CID | {a.get('ipfs_cid') or '—'} |",
    ]

    # ── Section 4: Graph Topology
    ntc = a.get("node_type_counts") or {}
    node_rows = [f"| {nt} | {cnt} |" for nt, cnt in ntc.items()] or ["| — | — |"]
    cp_rows = []
    for n in (a.get("top_counterparties") or [])[:10]:
        addr = n.get("address") or n.get("id") or "—"
        nt = n.get("node_type", "—")
        txc = n.get("tx_count") or 0
        risk_s = n.get("risk_score")
        risk_s_txt = f"{risk_s:.2f}" if risk_s is not None else "—"
        cp_rows.append(f"| `{addr[:20]}…` | {nt} | {txc} | {risk_s_txt} |")
    if not cp_rows:
        cp_rows = ["| — | — | — | — |"]

    # ── Section 5: Threat Intel
    threats = a.get("threats") or []
    threat_rows = []
    for t in threats:
        src = t.get("source") or "—"
        cat = t.get("category") or "—"
        conf = t.get("confidence")
        conf_txt = f"{conf:.0%}" if conf is not None else "—"
        threat_rows.append(f"| {src} | {cat} | {conf_txt} |")
    if not threat_rows:
        threat_rows = ["| No threat intelligence records found | — | — |"]

    # ── Section 6: Anomaly detection
    triggered_codes = {str(f.code) for f in (risk.anomaly_flags or [])}
    anom_rows = []
    for code, name in _ANOMALY_DEFS:
        status = "● TRIGGERED" if code in triggered_codes else "○ clear"
        sev = "—"
        desc = "—"
        for f in (risk.anomaly_flags or []):
            if str(f.code) == code:
                sev = f"{f.severity * 100:.0f}%"
                desc = f.description
        anom_rows.append(f"| `{code}` | {name} | {status} | {sev} | {desc} |")

    # ── Section 7: GNN
    gnn_info = a.get("gnn") or {}
    shap_rows = ""
    if risk.gnn and risk.gnn.shap_top_features:
        shap_rows = "\n".join(
            f"| {i+1} | {f['feature']} | {f['value']:.4f} | {f.get('raw_value', '—')} |"
            for i, f in enumerate(risk.gnn.shap_top_features[:10])
        )
    elif gnn_info.get("shap_top_features"):
        shap_rows = "\n".join(
            f"| {i+1} | {f['feature']} | {f['value']:.4f} | {f.get('raw_value', '—')} |"
            for i, f in enumerate(gnn_info["shap_top_features"][:10])
        )

    # Pre-compute GNN fields to avoid invalid f-string format specs
    gnn_score_txt = f"{risk.gnn.score:.4f}" if risk.gnn else "—"
    gnn_label_txt = risk.gnn.label.value if risk.gnn else "—"
    gnn_subgraph_txt = str(
        gnn_info.get("subgraph_size") or
        (getattr(risk.gnn, "subgraph_size", None) if risk.gnn else None) or "—"
    )

    md = f"""# TAHRIX Forensic Investigation Report

**Case Reference:** `{case_id}`  
**Generated:** {now} UTC  
**Classification:** RESTRICTED

---

## 9 · Agent Verdict: {grade} RISK — {score:.1f} / 100

> {expl}

---

## 1 · Subject of Investigation

| Field | Value |
|-------|-------|
| Target Address | `{address}` |
| Blockchain | {chain} |
| Case Reference | `{case_id}` |
| Generated | {now} UTC |
| OFAC Sanctions | {sanctions_txt} |
| GNN Illicit Probability | {gnn_txt} |

---

## 2 · Risk Scoring Breakdown

| Factor | Raw Score | Weight | Contribution |
|--------|-----------|--------|--------------|
{chr(10).join(comp_rows)}

---

## 3 · On-Chain Activity

| Metric | Value |
|--------|-------|
{chr(10).join(onchain_rows)}

---

## 4 · Knowledge Graph Topology

### Node Type Distribution

| Node Type | Count |
|-----------|-------|
{chr(10).join(node_rows)}

**Total Nodes:** {a.get('graph_node_count', 0)}  **Total Edges:** {a.get('graph_edge_count', 0)}

### Top Counterparties (by tx volume)

| Address | Type | Tx Count | Risk Score |
|---------|------|----------|------------|
{chr(10).join(cp_rows)}

---

## 5 · Threat Intelligence

| Source | Category | Confidence |
|--------|----------|------------|
{chr(10).join(threat_rows)}

---

## 6 · Anomaly Pattern Detection (P01–P17)

| Code | Pattern | Status | Severity | Description |
|------|---------|--------|----------|-------------|
{chr(10).join(anom_rows)}

---

## 7 · Graph Neural Network Analysis

| Field | Value |
|-------|-------|
| Model | GAT (Graph Attention Network) — Elliptic AUC 0.9684 |
| Illicit Probability | {gnn_score_txt} |
| Label | {gnn_label_txt} |
| Subgraph Size | {gnn_subgraph_txt} nodes |

{"### SHAP Feature Attribution" if shap_rows else ""}
{"| # | Feature | SHAP Value | Raw Value |" if shap_rows else ""}
{"|---|---------|-----------|-----------|" if shap_rows else ""}
{shap_rows}

---

## 8 · Investigation Execution

| Field | Value |
|-------|-------|
| Cognitive Iterations | {a.get('iterations', 0)} |
| Transactions Collected | {a.get('transactions_collected', 0)} |
| Duration | {a.get('duration_s') or '—'} s |
| Tools Invoked | {a.get('tool_call_count', '—')} |

---

*TAHRIX Agentic AI Cyber Intelligence Platform — Risk scores are advisory; verify with a qualified analyst before any legal action.*
"""
    return md.encode("utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────
def build_docx_report(*, case_id: str, risk, agent_result: dict[str, Any],
                      graph_svg: str | None = None) -> bytes:
    """Generate a structured DOCX report via python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ctx = _build_context(case_id=case_id, risk=risk, agent_result=agent_result)
    grade = (risk.grade or "low").upper()
    score = risk.score or 0
    now = ctx["generated_at"]
    address = risk.address
    chain = ctx["chain"]

    GRADE_COLORS = {
        "CRITICAL": RGBColor(0xDC, 0x26, 0x26),
        "HIGH":     RGBColor(0xEA, 0x58, 0x0C),
        "MEDIUM":   RGBColor(0xCA, 0x8A, 0x04),
        "LOW":      RGBColor(0x16, 0xA3, 0x4A),
    }
    grade_color = GRADE_COLORS.get(grade, RGBColor(0x0F, 0x17, 0x2A))

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    def _heading(text: str, level: int = 1):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    def _kv_table(rows: list[tuple[str, str]]):
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (k, v) in enumerate(rows):
            t.cell(i, 0).text = k
            t.cell(i, 0).paragraphs[0].runs[0].bold = True
            t.cell(i, 1).text = v
        return t

    def _section(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    # ── Title block
    p = doc.add_paragraph()
    r = p.add_run("TAHRIX")
    r.bold = True; r.font.size = Pt(22)
    doc.add_paragraph("Agentic Blockchain Intelligence · Forensic Report").runs[0].font.size = Pt(8)
    doc.add_paragraph(f"Case: {case_id}  |  Generated: {now} UTC  |  Classification: RESTRICTED").runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    # ── Verdict
    p = doc.add_paragraph()
    r = p.add_run(f"{grade} RISK — {score:.1f} / 100")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = grade_color
    if risk.explanation:
        doc.add_paragraph(risk.explanation).runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph()

    # ── Section 1
    _section("1 · Subject of Investigation")
    sanctions_txt = "SANCTIONED — OFAC SDN LIST MATCH" if (
        risk.sanctions and risk.sanctions.sanctioned) else "No match found"
    gnn_txt = (f"{risk.gnn.score:.4f} ({risk.gnn.label.value})" if risk.gnn else "Unavailable")
    _kv_table([
        ("Target Address", address),
        ("Blockchain Network", chain),
        ("Case Reference", case_id),
        ("Generated", f"{now} UTC"),
        ("OFAC Sanctions", sanctions_txt),
        ("GNN Illicit Probability", gnn_txt),
    ])

    doc.add_paragraph()

    # ── Section 2: Risk components
    _section("2 · Risk Scoring Breakdown")
    c = risk.components
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(["Factor", "Raw Score", "Weight", "Contribution"]):
        t.cell(0, i).text = h
        t.cell(0, i).paragraphs[0].runs[0].bold = True
    data = [
        ("GNN (Graph Neural Network)", f"{c.get('gnn', 0):.3f}", f"{c.get('w_gnn', 0)*100:.0f}%",
         f"{c.get('gnn', 0)*c.get('w_gnn', 0)*100:.1f}"),
        ("Anomaly Patterns", f"{c.get('anomaly', 0):.3f}", f"{c.get('w_anomaly', 0)*100:.0f}%",
         f"{c.get('anomaly', 0)*c.get('w_anomaly', 0)*100:.1f}"),
        ("OFAC Sanctions", f"{c.get('sanctions', 0):.3f}", f"{c.get('w_sanctions', 0)*100:.0f}%",
         f"{c.get('sanctions', 0)*c.get('w_sanctions', 0)*100:.1f}"),
        ("Threat Intelligence", f"{c.get('threat', 0):.3f}", f"{c.get('w_threat', 0)*100:.0f}%",
         f"{c.get('threat', 0)*c.get('w_threat', 0)*100:.1f}"),
        ("Network Centrality", f"{c.get('centrality', 0):.3f}", f"{c.get('w_centrality', 0)*100:.0f}%",
         f"{c.get('centrality', 0)*c.get('w_centrality', 0)*100:.1f}"),
        ("COMPOSITE SCORE", f"{score:.1f} / 100", "", ""),
    ]
    for row in data:
        r = t.add_row()
        for j, val in enumerate(row):
            r.cells[j].text = val
        if row[0] == "COMPOSITE SCORE":
            for j in range(4):
                for run in r.cells[j].paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()

    # ── Section 3: On-Chain Activity
    doc.add_paragraph()
    _section("3 · On-Chain Activity")
    tx_info = agent_result.get("tx_info") or {}
    _kv_table([
        ("Transactions Analysed", str(tx_info.get("count", agent_result.get("transactions_collected", 0)))),
        ("First Seen", str(tx_info.get("first_seen") or "—")),
        ("Last Seen", str(tx_info.get("last_seen") or "—")),
        ("Forward Trace Hops", str(agent_result.get("trace_fwd_count", 0))),
        ("Backward Trace Hops", str(agent_result.get("trace_bwd_count", 0))),
        ("Bridge Events", str(len(agent_result.get("bridge_events", [])))),
        ("IPFS Archival CID", str(agent_result.get("ipfs_cid") or "—")),
    ])

    # ── Section 4: Graph Topology
    doc.add_paragraph()
    _section("4 · Knowledge Graph Topology")
    ntc = agent_result.get("node_type_counts") or {}
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    for i, h in enumerate(["Node Type", "Count"]):
        t.cell(0, i).text = h
        t.cell(0, i).paragraphs[0].runs[0].bold = True
    for nt, cnt in ntc.items():
        r = t.add_row()
        r.cells[0].text = nt
        r.cells[1].text = str(cnt)
    doc.add_paragraph(
        f"Total nodes: {agent_result.get('graph_node_count', 0)}  "
        f"Total edges: {agent_result.get('graph_edge_count', 0)}"
    ).runs[0].font.size = Pt(8)
    doc.add_paragraph()
    doc.add_paragraph("Top Counterparties").runs[0].bold = True
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    for i, h in enumerate(["Address", "Type", "Tx Count", "Risk Score"]):
        t2.cell(0, i).text = h
        t2.cell(0, i).paragraphs[0].runs[0].bold = True
    for n in (agent_result.get("top_counterparties") or [])[:10]:
        addr = (n.get("address") or n.get("id") or "—")[:20] + "…"
        nt = n.get("node_type", "—")
        txc = str(n.get("tx_count") or 0)
        rs = n.get("risk_score")
        rstr = f"{rs:.2f}" if rs is not None else "—"
        r2 = t2.add_row()
        r2.cells[0].text = addr
        r2.cells[1].text = nt
        r2.cells[2].text = txc
        r2.cells[3].text = rstr

    # ── Section 5: Threat Intelligence
    doc.add_paragraph()
    _section("5 · Threat Intelligence")
    threats = agent_result.get("threats") or []
    if threats:
        t3 = doc.add_table(rows=1, cols=3)
        t3.style = "Table Grid"
        for i, h in enumerate(["Source", "Category", "Confidence"]):
            t3.cell(0, i).text = h
            t3.cell(0, i).paragraphs[0].runs[0].bold = True
        for th in threats:
            r3 = t3.add_row()
            r3.cells[0].text = str(th.get("source") or "—")
            r3.cells[1].text = str(th.get("category") or "—")
            conf = th.get("confidence")
            r3.cells[2].text = f"{conf:.0%}" if conf is not None else "—"
    else:
        doc.add_paragraph("No threat intelligence records found.").runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # ── Section 6: Anomaly patterns (renumbered)
    doc.add_paragraph()
    _section("6 · Anomaly Pattern Detection (P01–P17)")
    triggered_codes = {str(f.code) for f in (risk.anomaly_flags or [])}
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(["Code", "Pattern", "Status", "Severity"]):
        t.cell(0, i).text = h
        t.cell(0, i).paragraphs[0].runs[0].bold = True
    for code, name in _ANOMALY_DEFS:
        triggered = code in triggered_codes
        sev = "—"
        for f in (risk.anomaly_flags or []):
            if str(f.code) == code:
                sev = f"{f.severity*100:.0f}%"
        r = t.add_row()
        r.cells[0].text = code
        r.cells[1].text = name
        r.cells[2].text = "TRIGGERED" if triggered else "clear"
        r.cells[3].text = sev
        if triggered:
            for j in range(4):
                for run in r.cells[j].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    # ── Section 7: GNN + SHAP
    doc.add_paragraph()
    _section("7 · Graph Neural Network Analysis")
    gnn_info = agent_result.get("gnn") or {}
    _kv_table([
        ("Model", "GAT (Graph Attention Network) — Elliptic AUC 0.9684"),
        ("Illicit Probability", f"{risk.gnn.score:.4f}" if risk.gnn else "—"),
        ("Label", risk.gnn.label.value if risk.gnn else "—"),
        ("Subgraph Size", str(gnn_info.get("subgraph_size") or "—") + " nodes"),
    ])
    shap_feats = (risk.gnn.shap_top_features if risk.gnn and risk.gnn.shap_top_features
                  else gnn_info.get("shap_top_features") or [])
    if shap_feats:
        doc.add_paragraph()
        doc.add_paragraph("SHAP Feature Attribution").runs[0].bold = True
        ts = doc.add_table(rows=1, cols=4)
        ts.style = "Table Grid"
        for i, h in enumerate(["#", "Feature", "SHAP Value", "Raw Value"]):
            ts.cell(0, i).text = h
            ts.cell(0, i).paragraphs[0].runs[0].bold = True
        for i, feat in enumerate(shap_feats[:10]):
            rs = ts.add_row()
            rs.cells[0].text = str(i + 1)
            rs.cells[1].text = feat.get("feature", "")
            rs.cells[2].text = f"{feat.get('value', 0):.4f}"
            rs.cells[3].text = str(feat.get("raw_value", "—"))

    # ── Section 8: Execution summary (renumbered)
    doc.add_paragraph()
    _section("8 · Investigation Execution")
    _kv_table([
        ("Cognitive Iterations", str(agent_result.get("iterations", 0))),
        ("Transactions Collected", str(agent_result.get("transactions_collected", 0))),
        ("Duration", f"{agent_result.get('duration_s') or '—'} s"),
        ("Tools Invoked", str(agent_result.get("tool_call_count", "—"))),
    ])

    # ── Section 9: Transaction Graph (SVG → PNG for DOCX)
    if graph_svg:
        doc.add_paragraph()
        _section("9 · Transaction Graph")
        doc.add_paragraph("Connected nodes only. Focal address highlighted in blue. Isolated nodes omitted.").runs[0].font.size = Pt(8)
        try:
            import cairosvg
            png_bytes = cairosvg.svg2png(bytestring=graph_svg.encode(), scale=1.5)
            doc.add_picture(io.BytesIO(png_bytes), width=Cm(16))
        except Exception:
            # cairosvg unavailable — embed as base64 data-URI alt text note
            doc.add_paragraph("[Graph image: export as PDF to view the embedded graph]")

    # ── Footer
    doc.add_paragraph()
    p = doc.add_paragraph("TAHRIX Agentic AI Cyber Intelligence Platform — Risk scores are advisory; verify with a qualified analyst before any legal action.")
    p.runs[0].font.size = Pt(7)
    p.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def write_pdf_to_file(path: str | Path, *, case_id: str, risk,
                      agent_result: dict[str, Any]) -> Path:
    p = Path(path)
    p.write_bytes(build_pdf_report(case_id=case_id, risk=risk, agent_result=agent_result))
    return p
